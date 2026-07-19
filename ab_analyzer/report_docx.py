"""
Gera o relatório do teste A/B em Word (.docx), formatado para apresentação a gestor.

Reaproveita os mesmos dados da análise (métricas, decisão, outliers) e produz um
documento com título, banner de decisão colorido, tabela de métricas com a
variante vencedora destacada, e as seções de significância, trade-off, outliers,
qualidade de dados e metodologia.
"""
from __future__ import annotations

from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .decision import Decision, _fmt_brl, _fmt_pct, fmt_p
from .loader import Dataset

# paleta
INK = RGBColor(0x1F, 0x38, 0x64)      # azul-escuro (títulos)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_TX = RGBColor(0x1E, 0x6B, 0x2E)
AMBER_TX = RGBColor(0x8A, 0x63, 0x00)
GREY_TX = RGBColor(0x44, 0x44, 0x44)

HEADER_FILL = "1F3864"   # cabeçalho da tabela
WIN_FILL = "E6F4EA"      # linha da vencedora
GREEN_FILL = "E6F4EA"
AMBER_FILL = "FEF3CD"
GREY_FILL = "EDEDED"

BADGE = {
    "escalar": ("✅ ESCALAR", GREEN_FILL, GREEN_TX),
    "escalar_com_cautela": ("🟡 ESCALAR COM CAUTELA", AMBER_FILL, AMBER_TX),
    "inconclusivo": ("⚪ INCONCLUSIVO", GREY_FILL, GREY_TX),
}


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _cell(cell, text, *, bold=False, color=None, align=None, size=None, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.size = Pt(size if size else 10)
    if fill:
        _shade(cell, fill)


def _heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = INK
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def build_docx(ds: Dataset, metrics: pd.DataFrame, decision: Decision,
               outliers, test_name: str, description: str, out_path: str) -> str:
    doc = Document()
    # fonte padrão
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # ---- Título ----
    t = doc.add_paragraph()
    r = t.add_run(f"Relatório de Teste A/B — {test_name}")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = INK

    sub = doc.add_paragraph()
    rs = sub.add_run("Méliuz · Operações Integradas — análise de cashback")
    rs.italic = True
    rs.font.size = Pt(10)
    rs.font.color.rgb = GREY_TX

    meta = doc.add_paragraph()
    meta.add_run(
        f"Parceiro: {ds.partner}   |   Período: {ds.date_min} a {ds.date_max}   |   "
        f"Variantes: {ds.n_variants}   |   Observações: {ds.n_rows} linhas"
    ).font.size = Pt(9)
    if description:
        d = doc.add_paragraph()
        d.add_run(f"Descrição: {description}").font.size = Pt(9)

    # ---- Banner de decisão ----
    label, fill, txt_color = BADGE[decision.recommendation]
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    bcell = banner.rows[0].cells[0]
    _shade(bcell, fill)
    p1 = bcell.paragraphs[0]
    run1 = p1.add_run(f"{label} — {decision.winner}")
    run1.bold = True
    run1.font.size = Pt(15)
    run1.font.color.rgb = txt_color
    p2 = bcell.add_paragraph()
    run2 = p2.add_run(decision.headline)
    run2.font.size = Pt(10)
    run2.font.color.rgb = GREY_TX

    # ---- Tabela de métricas ----
    _heading(doc, "Métricas por variante")
    note = doc.add_paragraph()
    note.add_run("Métrica de decisão (OEC): margem líquida = comissão − cashback. "
                 "As demais são métricas-guardrail (contexto).").italic = True
    note.runs[0].font.size = Pt(8)

    cols = ["Variante", "Dias", "Compradores", "GMV", "Comissão", "Cashback",
            "Margem líquida", "Ticket", "% Cashback", "Margem/compr."]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(cols):
        _cell(table.rows[0].cells[j], c, bold=True, color=WHITE, fill=HEADER_FILL,
              align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

    R = WD_ALIGN_PARAGRAPH.RIGHT
    for v, row in metrics.iterrows():
        cells = table.add_row().cells
        is_win = (v == decision.winner)
        vals = [
            (f"{v}" + (" 🏆" if is_win else ""), WD_ALIGN_PARAGRAPH.LEFT),
            (f"{int(row['dias'])}", R),
            (f"{int(row['compradores']):,}".replace(",", "."), R),
            (_fmt_brl(row["gmv"]), R),
            (_fmt_brl(row["comissao"]), R),
            (_fmt_brl(row["cashback"]), R),
            (_fmt_brl(row["margem_liquida"]), R),
            (_fmt_brl(row["ticket_medio"]), R),
            (_fmt_pct(row["cashback_pct_gmv"]), R),
            (_fmt_brl(row["margem_por_comprador"]), R),
        ]
        for j, (txt, al) in enumerate(vals):
            bold = is_win and j in (0, 6)
            _cell(cells[j], txt, bold=bold, align=al, size=9,
                  fill=WIN_FILL if is_win else None)

    # ---- Por que essa variante ----
    _heading(doc, "Por que essa variante")
    for line in decision.rationale:
        _bullet(doc, line)

    # ---- Significância ----
    t = decision.test_result
    _heading(doc, "Significância estatística")
    doc.add_paragraph(
        f"Comparação da margem líquida diária entre a vencedora ({t.variant_a}) e a "
        f"2ª colocada ({t.variant_b}), via teste de permutação (p-valor) e bootstrap (IC)."
    ).runs[0].font.size = Pt(10)
    _bullet(doc, f"Diferença de média diária: {_fmt_brl(t.diff)}/dia")
    _bullet(doc, f"Intervalo de confiança 95%: {_fmt_brl(t.ci_low)} a {_fmt_brl(t.ci_high)}")
    _bullet(doc, f"p-valor (bicaudal): {fmt_p(t.p_value)} "
                 f"({'significativo' if t.significant else 'não significativo'} a 5%)")
    _bullet(doc, f"Tamanho de efeito (d de Cohen): {t.cohens_d:.2f}  |  Amostra: {t.n_a} vs {t.n_b} dias")

    # ---- Trade-off ----
    _heading(doc, "Trade-off (leitura crítica)")
    for line in decision.tradeoff:
        _bullet(doc, line)

    # ---- Outliers ----
    _heading(doc, "Outliers e robustez da decisão")
    if outliers.total_outliers == 0:
        _bullet(doc, "Nenhum outlier detectado na margem líquida diária (método IQR).")
    else:
        _bullet(doc, f"{outliers.total_outliers} dia(s) atípico(s) detectado(s) "
                     f"(método IQR, 1,5×):")
        for v, info in outliers.per_variant.items():
            if info["outliers"]:
                dates = ", ".join(o["date"] for o in info["outliers"])
                _bullet(doc, f"{v}: {len(info['outliers'])} dia(s) — {dates}")
        if outliers.concurrent_dates:
            _bullet(doc, f"Datas atípicas concorrentes (≥2 variantes no mesmo dia): "
                         f"{', '.join(outliers.concurrent_dates)} — padrão de "
                         f"evento/promoção real, não erro isolado; baixo impacto no A/B.")
    robust = "robusta" if outliers.is_robust else "sensível"
    _bullet(doc, f"Teste de robustez (winsorização 5%/95%): a vencedora se mantém "
                 f"{outliers.robust_winner} → decisão {robust} a outliers.")

    # ---- Qualidade dos dados ----
    _heading(doc, "Qualidade dos dados")
    for iss in ds.issues:
        _bullet(doc, f"[{iss.severity}] {iss.message}")

    # ---- Metodologia ----
    _heading(doc, "Metodologia e ressalvas")
    _bullet(doc, "Decisão por margem líquida: escalar uma variante é decisão de P&L "
                 "(o Méliuz ganha comissão e paga cashback).")
    _bullet(doc, "Ressalva: os dados têm compradores (conversões), mas não sessões/"
                 "visitantes por variante — assume-se split de tráfego equilibrado.")
    _bullet(doc, "Significância por reamostragem (permutação + bootstrap), sem suposição "
                 "de normalidade.")

    foot = doc.add_paragraph()
    fr = foot.add_run(f"Gerado automaticamente em {datetime.now():%d/%m/%Y %H:%M} "
                      f"pela solução de análise de testes A/B.")
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = GREY_TX

    doc.save(out_path)
    return out_path
